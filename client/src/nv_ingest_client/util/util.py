# SPDX-FileCopyrightText: Copyright (c) 2024, NVIDIA CORPORATION & AFFILIATES.
# All rights reserved.
# SPDX-License-Identifier: Apache-2.0

import glob
import logging
import os
import traceback
import typing
from io import BytesIO
from typing import Dict

import pypdfium2 as pdfium
from docx import Document as DocxDocument
from nv_ingest_client.util.file_processing.extract import DocumentTypeEnum
from nv_ingest_client.util.file_processing.extract import detect_encoding_and_read_text_file
from nv_ingest_client.util.file_processing.extract import extract_file_content
from nv_ingest_client.util.file_processing.extract import get_or_infer_file_type
from pptx import Presentation

# pylint: disable=invalid-name
# pylint: disable=missing-class-docstring
# pylint: disable=logging-fstring-interpolation

logger = logging.getLogger(__name__)


def estimate_page_count(file_path: str) -> int:
    document_type = get_or_infer_file_type(file_path)

    if document_type in [
        DocumentTypeEnum.pdf,
        DocumentTypeEnum.docx,
        DocumentTypeEnum.pptx,
    ]:
        return count_pages_for_documents(file_path, document_type)
    elif document_type in [
        DocumentTypeEnum.txt,
        DocumentTypeEnum.md,
        DocumentTypeEnum.html,
    ]:
        return count_pages_for_text(file_path)
    elif document_type in [
        DocumentTypeEnum.jpeg,
        DocumentTypeEnum.bmp,
        DocumentTypeEnum.png,
        DocumentTypeEnum.svg,
    ]:
        return 1  # Image types assumed to be 1 page
    else:
        return 0


def count_pages_for_documents(file_path: str, document_type: DocumentTypeEnum) -> int:
    try:
        if document_type == DocumentTypeEnum.pdf:
            doc = pdfium.PdfDocument(file_path)
            return len(doc)
        elif document_type == DocumentTypeEnum.docx:
            doc = DocxDocument(file_path)
            # Approximation, as word documents do not have a direct 'page count' attribute
            return len(doc.paragraphs) // 15
        elif document_type == DocumentTypeEnum.pptx:
            ppt = Presentation(file_path)
            return len(ppt.slides)
    except FileNotFoundError:
        print(f"The file {file_path} was not found.")
        return 0
    except Exception as e:
        print(f"An error occurred while processing {file_path}: {e}")
        return 0


def count_pages_for_text(file_path: str) -> int:
    """
    Estimates the page count for text files based on word count,
    using the detect_encoding_and_read_text_file function for reading.
    """
    try:
        with open(file_path, "rb") as file:  # Open file in binary mode
            file_stream = BytesIO(file.read())  # Create BytesIO object from file content

        content = detect_encoding_and_read_text_file(file_stream)  # Read and decode content
        word_count = len(content.split())
        pages_estimated = word_count / 300
        return round(pages_estimated)
    except FileNotFoundError:
        logger.error(f"The file {file_path} was not found.")
        return 0
    except Exception as e:
        logger.error(f"An error occurred while processing {file_path}: {e}")
        return 0


def _process_file(file_path: str):
    """
    Synchronously processes a single file, extracting its content and collecting file details.

    This function serves as a high-level interface for file processing, invoking content
    extraction and aggregating the results along with file metadata. It is designed to work
    within a larger processing pipeline, providing necessary data for subsequent tasks or
    storage.

    Parameters
    ----------
    file_path : str
        The path to the file that needs to be processed.

    Returns
    -------
    dict
        A dictionary containing details about the processed file, including its name, a unique
        identifier, the extracted content, and the document type.

    Raises
    ------
    Exception
        Propagates any exceptions encountered during the file processing, signaling issues with
        content extraction or file handling.

    Notes
    -----
    - The function directly utilizes `extract_file_content` for content extraction and performs
      basic error handling.
    - It constructs a simple metadata object that can be utilized for further processing or
      logging.
    """

    try:
        file_name = os.path.basename(file_path)
        content, document_type = extract_file_content(file_path)  # Call the synchronous function directly

        return {
            "source_name": file_name,
            "source_id": file_name,
            "content": content,
            "document_type": document_type,
        }
    except Exception as e:
        traceback.print_exc()
        logger.error(f"Error processing file {file_path}: {e}")
        raise


def load_data_from_path(path: str) -> Dict:
    """
    Loads data from a specified file path, preparing it for processing.

    Parameters
    ----------
    path : str
        The path to the file from which data should be loaded.

    Returns
    -------
    dict
        A dictionary containing keys 'file_name', 'id', 'content', and 'document_type',
        each of which maps to a list that includes the respective details for the processed file.

    Raises
    ------
    FileNotFoundError
        If the specified path does not exist.
    ValueError
        If the specified path is not a file.

    Notes
    -----
    This function is designed to load and prepare file data for further processing,
    packaging the loaded data along with metadata such as file name and document type.
    """

    result = {"source_name": [], "source_id": [], "content": [], "document_type": []}

    if not os.path.exists(path):
        raise FileNotFoundError(f"The path {path} does not exist.")

    if not os.path.isfile(path):
        raise ValueError("The provided path is not a file.")

    file_data = _process_file(file_path=path)
    result["content"].append(file_data["content"])
    result["document_type"].append(file_data["document_type"])
    result["source_name"].append(file_data["source_name"])
    result["source_id"].append(file_data["source_id"])

    return result


def check_ingest_result(json_payload: Dict) -> typing.Tuple[bool, str]:
    """
    Check the ingest result to determine if the process failed and extract a description.

    This function examines the provided JSON payload to check whether the ingest operation
    has failed. If it has failed and failure annotations are present, the function augments
    the failure description with information about the specific event that caused the failure.

    Parameters
    ----------
    json_payload : Dict
        A dictionary containing the result of the ingest operation. The dictionary should
        have at least the following fields:
        - 'status' (str): The status of the ingest operation. A status of "failed" indicates
          a failure.
        - 'description' (str): A textual description of the result.
        - 'annotations' (Dict): (optional) A dictionary of annotations, where each annotation
          may contain details about the failure.

    Returns
    -------
    Tuple[bool, str]
        A tuple containing:
        - A boolean indicating whether the ingest operation failed (`True` if it failed,
          `False` otherwise).
        - A string containing the description of the result or the failure, augmented with
          details from the annotations if available.

    Notes
    -----
    - The function logs the 'status' and 'description' fields from the payload for debugging.
    - If the 'status' field contains "failed" and the 'annotations' field contains entries
      indicating failure, the function updates the description with the annotation details.
    - The function breaks after processing the first relevant annotation.

    Examples
    --------
    Suppose the JSON payload contains the following structure:

    >>> json_payload = {
    ...     "status": "failed",
    ...     "description": "Ingest operation failed",
    ...     "annotations": {
    ...         "task1": {"task_result": "FAILURE", "message": "Network error"},
    ...         "task2": {"task_result": "SUCCESS"}
    ...     }
    ... }
    >>> is_failed, description = check_ingest_result(json_payload)
    >>> print(is_failed)
    True
    >>> print(description)
    ↪ Event that caused this failure: task1 -> Network error

    In this example, the function detects a failure and augments the description with the
    message from the failing task.
    """

    logger.debug(
        f"Checking ingest result:\n Status: {json_payload.get('status', None)}"
        f"\n Description: {json_payload.get('description', None)}")

    is_failed = json_payload.get("status", "") in "failed"
    description = json_payload.get("description", "")

    # Look to see if we have any failure annotations to augment the description
    if (is_failed and 'annotations' in json_payload):
        for annot_id, value in json_payload['annotations'].items():
            if ('task_result' in value and value['task_result'] == "FAILURE"):
                message = value.get('message', "Unknown")
                description = f"\n↪ Event that caused this failure: {annot_id} -> {message}"
                break

    return is_failed, description


def generate_matching_files(file_sources):
    """
    Generates a list of file paths that match the given patterns specified in file_sources.

    Parameters
    ----------
    file_sources : list of str
        A list containing the file source patterns to match against.

    Returns
    -------
    generator
        A generator yielding paths to files that match the specified patterns.

    Notes
    -----
    This function utilizes glob pattern matching to find files that match the specified patterns.
    It yields each matching file path, allowing for efficient processing of potentially large
    sets of files.
    """

    files = [
        file_path
        for pattern in file_sources
        for file_path in glob.glob(pattern, recursive=True)
        if os.path.isfile(file_path)
    ]
    for file_path in files:
        yield file_path
