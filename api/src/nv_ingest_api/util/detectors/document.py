# SPDX-FileCopyrightText: Copyright (c) 2024, NVIDIA CORPORATION & AFFILIATES.
# All rights reserved.
# SPDX-License-Identifier: Apache-2.0


# pylint: disable=invalid-name
# pylint: disable=missing-class-docstring
# pylint: disable=logging-fstring-interpolation

import base64
from io import BytesIO
import logging
import os
import charset_normalizer
from typing import Tuple

import pypdfium2 as pdfium
from docx import Document as DocxDocument
from pptx import Presentation

from nv_ingest_api.internal.enums.common import DocumentTypeEnum

logger = logging.getLogger(__name__)


# Maps file extensions to DocumentTypeEnum
EXTENSION_TO_DOCUMENT_TYPE = {
    "bmp": DocumentTypeEnum.BMP,
    "docx": DocumentTypeEnum.DOCX,
    "html": DocumentTypeEnum.HTML,
    "jpeg": DocumentTypeEnum.JPEG,
    "jpg": DocumentTypeEnum.JPEG,
    "json": DocumentTypeEnum.TXT,
    "md": DocumentTypeEnum.TXT,
    "pdf": DocumentTypeEnum.PDF,
    "png": DocumentTypeEnum.PNG,
    "pptx": DocumentTypeEnum.PPTX,
    "sh": DocumentTypeEnum.TXT,
    "svg": DocumentTypeEnum.SVG,
    "tiff": DocumentTypeEnum.TIFF,
    "txt": DocumentTypeEnum.TXT,
    "mp3": DocumentTypeEnum.MP3,
    "wav": DocumentTypeEnum.WAV,
    # Add more as needed
}


def serialize_to_base64(file_stream: BytesIO) -> str:
    """Reads a PDF file from a BytesIO object and encodes it in base64."""
    try:
        content = base64.b64encode(file_stream.read()).decode("utf-8")
        return content
    except IOError:
        logger.error("Failed to read PDF file from BytesIO object")
        raise


def extract_file_content(path: str) -> Tuple[str, DocumentTypeEnum]:
    """Extracts content from a file, supporting different formats."""
    document_type = get_or_infer_file_type(path)

    with open(path, "rb") as file:
        file_stream = BytesIO(file.read())

    try:
        if document_type in [
            DocumentTypeEnum.TXT,
            DocumentTypeEnum.MD,
            DocumentTypeEnum.HTML,
        ]:
            content = detect_encoding_and_read_text_file(file_stream)
        else:
            content = serialize_to_base64(file_stream)
    except Exception as e:
        logger.error(f"Error processing file {path}: {e}")

        raise ValueError(f"Failed to extract content from {path}") from e

    logger.debug(f"Content extracted from '{path}'")
    return content, DocumentTypeEnum(document_type)


def get_or_infer_file_type(file_path: str) -> DocumentTypeEnum:
    """
    Determines the file type by inspecting its extension and optionally falls back
    to MIME type detection if the extension is not recognized.

    Parameters
    ----------
    file_path : str
        The path to the file.

    Returns
    -------
    DocumentTypeEnum
        An enum value representing the detected file type.

    Raises
    ------
    ValueError
        If a valid extension is not found and MIME type detection cannot determine a valid type.
    """
    extension = os.path.splitext(file_path)[1][1:].lower()
    file_type = EXTENSION_TO_DOCUMENT_TYPE.get(extension)

    # If the file extension maps to a known type, return it
    if file_type:
        return file_type

    # TODO(Devin): libmagic is missing on the CI system, so we need to skip this check
    # If extension is not recognized, attempt MIME type detection as a fallback
    # mime_type = magic.from_file(file_path, mime=True)
    # # Attempt to map MIME type to DocumentTypeEnum, if possible
    # for mime, doc_type in MIME_TO_DOCUMENT_TYPE.items():
    #     if mime_type == mime:
    #         return doc_type

    # If no valid file type is determined, raise an exception
    raise ValueError(f"Failed to determine file type for: {file_path}")


def detect_encoding_and_read_text_file(file_stream: BytesIO) -> str:
    """Detects encoding and reads a text file from a BytesIO object accordingly."""
    try:
        raw_data = file_stream.read()
        result = charset_normalizer.detect(raw_data)
        encoding = result.get("encoding", "utf-8")  # Fallback to utf-8 if undetected

        content = raw_data.decode(encoding)
        return content
    except IOError:
        logger.error("Failed to read text file from BytesIO object")
        raise


def estimate_page_count(file_path: str) -> int:
    document_type = get_or_infer_file_type(file_path)

    if document_type in [
        DocumentTypeEnum.PDF,
        DocumentTypeEnum.DOCX,
        DocumentTypeEnum.PPTX,
    ]:
        return count_pages_for_documents(file_path, document_type)
    elif document_type in [
        DocumentTypeEnum.TXT,
        DocumentTypeEnum.MD,
        DocumentTypeEnum.HTML,
    ]:
        return count_pages_for_text(file_path)
    elif document_type in [
        DocumentTypeEnum.JPEG,
        DocumentTypeEnum.BMP,
        DocumentTypeEnum.PNG,
        DocumentTypeEnum.SVG,
    ]:
        return 1  # Image types assumed to be 1 page
    else:
        return 0


def count_pages_for_documents(file_path: str, document_type: DocumentTypeEnum) -> int:
    try:
        if document_type == DocumentTypeEnum.PDF:
            doc = pdfium.PdfDocument(file_path)
            return len(doc)
        elif document_type == DocumentTypeEnum.DOCX:
            doc = DocxDocument(file_path)
            # Approximation, as word documents do not have a direct 'page count' attribute
            return len(doc.paragraphs) // 15
        elif document_type == DocumentTypeEnum.PPTX:
            ppt = Presentation(file_path)
            return len(ppt.slides)
    except FileNotFoundError:
        print(f"The file {file_path} was not found.")
        return 0
    except Exception as e:
        print(f"An error occurred while processing {file_path}: {e}")
        return 0


def count_pages_for_text(file_path: str) -> int:
    """
    Estimates the page count for text files based on word count,
    using the detect_encoding_and_read_text_file function for reading.
    """
    try:
        with open(file_path, "rb") as file:  # Open file in binary mode
            file_stream = BytesIO(file.read())  # Create BytesIO object from file content

        content = detect_encoding_and_read_text_file(file_stream)  # Read and decode content
        word_count = len(content.split())
        pages_estimated = word_count / 300
        return round(pages_estimated)
    except FileNotFoundError:
        logger.error(f"The file {file_path} was not found.")
        return 0
    except Exception as e:
        logger.error(f"An error occurred while processing {file_path}: {e}")
        return 0
